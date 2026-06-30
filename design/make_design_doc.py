#!/usr/bin/env python3
"""Generate the Stage-1 one-page technical design document (DOCX).

Run with the sibling venv that has python-docx:
    ../../eightfold-candidate-transformer/.venv/bin/python make_design_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FULL_NAME = "Rupesh Jha"
EMAIL = "rujha@linkedin.com"
OUT = f"{FULL_NAME.replace(' ', '_')}_{EMAIL}_Eightfold.docx"

ACCENT = RGBColor(0x1F, 0x3B, 0x73)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Tight one-page geometry.
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.45)
    s.left_margin = s.right_margin = Inches(0.6)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(8)
normal.paragraph_format.space_after = Pt(2)
normal.paragraph_format.space_before = Pt(0)


def heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = ACCENT
    return p


def body(text, bullet=False, after=2):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    if bullet:
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.13)
    _emit_runs(p, text)
    return p


def _emit_runs(p, text):
    # **bold** and `mono` inline markup.
    import re
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not tok:
            continue
        r = p.add_run()
        if tok.startswith("**") and tok.endswith("**"):
            r.text = tok[2:-2]; r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r.text = tok[1:-1]; r.font.name = "Consolas"; r.font.size = Pt(8)
        else:
            r.text = tok


def figure(path, caption, width_in=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(3)
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(7); r.font.color.rgb = GREY


# ---- Title -----------------------------------------------------------------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(0)
r = t.add_run("Veritas — Multi-Source Candidate Data Transformer")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
r = sub.add_run(f"Technical Design (Stage 1) · {FULL_NAME} · {EMAIL} · implemented in Java (JDK 17, zero dependencies)")
r.font.size = Pt(8)
r.font.color.rgb = GREY

body("**Thesis.** Never merge raw values directly. Every source emits typed, traced **Claims** into a shared "
     "**Evidence Ledger**; the canonical profile is a pure, **deterministic reduction** over that ledger. One "
     "decision delivers all three constraints — deterministic & explainable (same inputs → same output; every "
     "field traces to a source + method), robust (a bad source adds zero claims, never crashes), and scalable "
     "(streaming, O(claims) per candidate).")

# ---- Pipeline (diagram) ----------------------------------------------------
heading("1 · Pipeline — input → engine → output")
figure("flow_diagram_compact.png",
       "Figure 1 — Many messy sources → traced Claims in one Evidence Ledger → deterministic reduction into a "
       "canonical record → a config-driven projection into the requested output.")
body("Each adapter **normalizes inside a sandbox** (so merge sees only clean, comparable values, with the raw "
     "kept for the audit trail); **merge** reduces the ledger to one canonical record; **confidence** scores each "
     "field; **project + validate** apply the runtime config and check it against the requested schema.", after=3)

# ---- Schema ----------------------------------------------------------------
heading("2 · Canonical schema & normalized formats")
body("Internal record (then rendered to the default schema): `candidate_id, full_name, emails[], phones[], "
     "location{city,region,country}, links{linkedin,github,portfolio,other[]}, headline, years_experience, "
     "skills[{name,confidence,sources[]}], experience[{company,title,start,end,summary}], "
     "education[{institution,degree,field,end_year}], provenance[{field,source,method}], overall_confidence`.")
body("Normalization (each returns **null on failure — never a guess**): phones → **E.164**; dates → **YYYY-MM**; "
     "country → **ISO-3166 alpha-2**; skills → **canonical names** via an alias table "
     "(`golang→Go`, `k8s→Kubernetes`); emails → lower-cased & deduped.", after=3)

# ---- Merge -----------------------------------------------------------------
heading("3 · Merge / conflict-resolution policy & confidence")
body("**Match key:** one candidate per manifest, so claims already share identity; fields key on their canonical "
     "path. **Scalar winner (deterministic order):** (1) highest **effective trust** — a per-source base with "
     "**per-field overrides** (recruiter CSV 0.90 & ATS 0.85 outrank LinkedIn 0.75; GitHub 0.95 on skills but "
     "0.0 on phones); (2) most recent timestamp; (3) earliest manifest order; (4) lexically smallest value. "
     "Steps 3–4 guarantee byte-identical output every run.")
body("**Lists** union & dedup; **skills** fold case-insensitively (free-text “distributed systems” ≡ ATS "
     "“Distributed Systems”); **experience/education** dedup by (company,title)/(institution,degree), back-filling "
     "blanks only. **Confidence** = authority + agreement (decaying per corroborating source) − conflict − a "
     "normalize penalty, clamped to [0,1]; `overall_confidence` is a weighted mean favouring identity + "
     "reachability.", after=3)

# ---- Config ----------------------------------------------------------------
heading("4 · Runtime custom-output config (projection + validation)")
body("The canonical record (a type) and the **projection layer** (a total interpreter of a config document) are "
     "**cleanly separated**, so a new output shape needs **no engine change**. A config can **select** a field "
     "subset, **rename/remap** from a canonical path (`\"from\"`), set **per-field normalize**, **toggle** "
     "provenance/confidence, and pick a missing-value policy `null | omit | error`. Each value is **type-validated** "
     "against the requested schema before return.")
body("Path mini-language for `from`: `emails[0]`, `location.country`, `skills[].name` (pluck per element), "
     "`links.other`. Example: `{\"path\":\"phone\",\"from\":\"phones[0]\",\"normalize\":\"E164\"}`.", after=3)

# ---- Edge cases ------------------------------------------------------------
heading("5 · Edge cases handled & deliberately descoped")
body("**Corrupt / malformed source** (e.g. broken JSON): adapter sandbox → zero claims, run continues; surfaced "
     "in the `--explain` log.", bullet=True)
body("**Conflicting identity** across sources (“Ada B. Lovelace” vs “Ada Byron Lovelace”): "
     "trust-ranked winner, confidence reduced to reflect the disagreement.", bullet=True)
body("**Duplicate values, different shape** (emails in mixed case; phones as `(415) 555-0142` vs `+1 415…`): "
     "normalize-then-dedup folds them to one.", bullet=True)
body("**Missing field / unknown alias:** honest `null` (or `omit`/`error` per config) — never invented; niche "
     "skills kept verbatim with confidence reflecting single-source support.", bullet=True)
body("**Descoped under time pressure:** cross-candidate identity resolution; ML résumé parsing (regex + section "
     "headers instead, for determinism); learned trust weights; live network fetch (GitHub response captured for "
     "reproducibility).", bullet=True, after=0)

doc.save(OUT)
print("wrote", OUT)
