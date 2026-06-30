#!/usr/bin/env python3
"""Generate a BEGINNER-FRIENDLY guide to the Veritas project (DOCX).

Explains, in plain language a newcomer can follow:
  - the problem the project solves,
  - the big idea behind the solution,
  - how data flows through the system, and
  - every single class, one at a time, with an analogy.

Run with the sibling venv that has python-docx:
    ../../eightfold-candidate-transformer/.venv/bin/python make_beginner_guide.py
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(os.path.dirname(HERE), "Veritas_Beginner_Guide.docx")
DIAGRAM = os.path.join(HERE, "flow_diagram.png")

ACCENT = RGBColor(0x1F, 0x3B, 0x73)   # deep blue
ACCENT2 = RGBColor(0x2E, 0x5A, 0x88)  # mid blue
GREEN = RGBColor(0x1E, 0x6B, 0x3A)
GREY = RGBColor(0x55, 0x55, 0x55)
CODEBG = "F2F2F2"
NOTEBG = "EAF1F8"

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.6)
    s.left_margin = s.right_margin = Inches(0.7)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.05


def _runs(p, text, size=None):
    """Render **bold** and `code` inline markup inside a paragraph."""
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not tok:
            continue
        r = p.add_run()
        if size:
            r.font.size = Pt(size)
        if tok.startswith("**") and tok.endswith("**"):
            r.text = tok[2:-2]; r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r.text = tok[1:-1]; r.font.name = "Consolas"; r.font.size = Pt(size or 9.5)
        else:
            r.text = tok


def shade(paragraph, fill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = ACCENT


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = ACCENT2


def body(text, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    _runs(p, text)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    _runs(p, text)


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    _runs(p, text)


def code(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.15)
    shade(p, CODEBG)
    for i, line in enumerate(lines):
        r = p.add_run(("" if i == 0 else "\n") + line)
        r.font.name = "Consolas"; r.font.size = Pt(9)


def note(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    shade(p, NOTEBG)
    r = p.add_run(label + "  "); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = ACCENT
    _runs(p, text, size=9.5)


def figure(path, caption, width_in=6.9):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        r = c.paragraphs[0].add_run(htext); r.bold = True; r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            _runs(cells[i].paragraphs[0], str(val), size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def classcard(filename, one_liner, plain, analogy, knobs=None):
    """A standard card describing one class for a beginner."""
    h2(filename)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("In one line:  "); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GREEN
    _runs(p, one_liner, size=9.5)
    body(plain)
    if knobs:
        for k in knobs:
            bullet(k)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Analogy:  "); r.italic = True; r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
    rr = p.add_run(analogy); rr.italic = True; rr.font.size = Pt(9.5); rr.font.color.rgb = GREY


# ===========================================================================
# Title block
# ===========================================================================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(0)
r = t.add_run("Veritas — A Beginner's Guide")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ACCENT

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(1)
r = sub.add_run("Understanding the Multi-Source Candidate Data Transformer")
r.font.size = Pt(12); r.font.color.rgb = ACCENT2

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.paragraph_format.space_after = Pt(8)
r = sub2.add_run("Java · JDK 17 · zero third-party libraries  |  Written for readers new to the codebase")
r.font.size = Pt(9.5); r.font.color.rgb = GREY

body("This guide assumes you have never seen this project before. It explains **what problem** the project "
     "solves, **the one big idea** behind the solution, **how data flows** from messy input to clean output, and "
     "then walks through **every class, one at a time**, with a plain-language analogy for each. By the end you "
     "should be able to open any file and know why it exists and what it does.")

# ===========================================================================
h1("1 · What is this project?")
body("**Veritas takes messy information about a job candidate from many different places and produces one clean, "
     "trustworthy profile.** Imagine you are hiring someone named Ada. You have:")
bullet("a spreadsheet (CSV) exported from a recruiting tool,")
bullet("a chunk of data (JSON) from an Applicant Tracking System with oddly-named fields,")
bullet("their public GitHub page,")
bullet("their LinkedIn profile,")
bullet("a résumé saved as plain text, and")
bullet("free-form notes a recruiter typed.")
body("Each one spells her name a little differently, lists different phone numbers, and disagrees about her "
     "skills. One file is even **corrupted**. Veritas reads all of them and outputs a single profile that says, "
     "for every field: *here is the best value, here is how confident we are, and here is exactly which source it "
     "came from.*")

note("Why \"Veritas\"?",
     "\"Veritas\" is Latin for **truth**. The project's job is to find the most truthful single answer when "
     "sources disagree — and to be honest about how sure it is.")

# ===========================================================================
h1("2 · The problem, explained simply")
body("Combining data from many sources sounds easy (\"just merge them\"), but it is full of traps. The project "
     "must solve four hard things at once:")

table(["The challenge", "What it means in plain English"],
      [["**Conflicts**", "Three sources say the name is \"Ada Lovelace\", one says \"Ada B. Lovelace\", one says "
                          "\"Ada Byron Lovelace\". Which is right?"],
       ["**Messy formats**", "Phone numbers appear as `(415) 555-0142`, `+1 415 555 0142`, `415-555-0199`. The "
                             "country is \"USA\" in one place and \"United States\" in another."],
       ["**Broken sources**", "One file is invalid and would normally crash the program. The run must survive and "
                              "use the other sources."],
       ["**Trust & proof**", "We must record *where every value came from* and *how confident* we are, so a human "
                             "can audit any decision."]])

note("The golden rule",
     "**\"Honestly empty\" beats \"confidently wrong.\"**  A made-up phone number that looks real is dangerous — it "
     "could send a rejection to the wrong person. So whenever Veritas is unsure, it returns **nothing (null)** "
     "rather than guessing.")

# ===========================================================================
h1("3 · The big idea behind the solution")
body("Almost everything in this codebase follows from **one decision**:")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
r = p.add_run("Never merge raw values directly. Turn every source into small, traceable \"Claims\", "
              "then calculate the final answer from those Claims.")
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ACCENT

body("A **Claim** is a single tiny statement, like: *\"The recruiter CSV says the full name is 'Ada B. Lovelace'; "
     "I trust this source 0.90 out of 1 for names; I saw it on 2026-06-20.\"*  Every source produces a list of "
     "these Claims. Put all the Claims for one candidate together and you get the **Evidence Ledger** — a big "
     "list of every statement anyone made about this person.")
body("The final profile is then **calculated** from the ledger: for each field, look at all the competing Claims, "
     "pick a winner using clear rules, and compute a confidence score. Because this calculation only depends on "
     "the (unchanging) list of Claims, it is:")
bullet("**Deterministic** — the same inputs always give the exact same output (great for testing).")
bullet("**Explainable** — every value can point back to the Claim it came from (that is the provenance).")
bullet("**Robust** — a broken source simply contributes zero Claims; nothing crashes.")
bullet("**Scalable** — each candidate is handled independently, so you can process millions in parallel.")

figure(DIAGRAM,
       "Figure 1 — The journey of the data: many messy sources on the left become traced Claims in one Evidence "
       "Ledger, which is reduced into a clean canonical record, which a configurable projection turns into the "
       "exact output shape the caller asked for.")

# ===========================================================================
h1("4 · How the data flows (the seven stages)")
body("Veritas processes a candidate in seven steps. Keep this list in mind — the classes in Section 6 each "
     "handle one or two of these stages.")

table(["Stage", "What happens", "Done by"],
      [["1 · detect", "Read the manifest — the little file that lists which sources to load.", "`Pipeline`"],
       ["2 · extract", "Run the right reader (\"adapter\") for each source, safely. Each produces Claims.",
        "`adapters/*`"],
       ["3 · normalize", "Clean each value into a standard format (phones → E.164, \"USA\" → \"US\").",
        "`Normalize`"],
       ["4 · merge", "For each field, pick the winning Claim and combine list fields.", "`Merge`"],
       ["5 · confidence", "Score how sure we are about each field and the profile overall.", "`Merge`"],
       ["6 · project", "Reshape the output according to a runtime config (the \"twist\").", "`Projection`"],
       ["7 · validate", "Check the result has the right shape and types before returning.", "`Validate`"]])

note("Stages 3–5 are where the magic is",
     "Normalization makes values comparable, merging resolves conflicts, and confidence tells the caller how much "
     "to trust each value. Everything else is plumbing around these three.")

# ===========================================================================
h1("5 · The key words you need (mini-glossary)")
body("These five terms appear everywhere. Learn them once and the code reads easily.")
table(["Word", "Beginner-friendly meaning"],
      [["**Claim**", "One small, traceable statement from one source about one field (name, email, etc.)."],
       ["**Evidence Ledger**", "The full list of all Claims for a candidate — the raw material for the answer."],
       ["**Trust**", "A number from 0 to 1 saying how much we believe a source for a particular field."],
       ["**Confidence**", "A number from 0 to 1 in the *output* saying how sure we are about a final value."],
       ["**Provenance**", "The paper trail: which source and which method produced a value."],
       ["**Canonical record**", "The project's clean, standard internal profile — the same shape every time."],
       ["**Projection**", "Reshaping that standard profile into whatever output a caller's config requests."],
       ["**Normalize**", "Convert a messy value into one standard format (or null if impossible)."]])

# ===========================================================================
h1("6 · Every class, explained one at a time")
body("The code lives under `src/main/java/com/eightfold/veritas/`. We will go package by package. For each class "
     "you get a one-line summary, a plain explanation, and an everyday analogy.")

# --- The model package ---
h2("Package: model/  — the data shapes")
body("These three classes are pure **data containers** — they hold information but contain little logic. They "
     "define the vocabulary the rest of the engine speaks.")

classcard("model/Claim.java",
          "One traceable statement from one source about one field.",
          "This is the atom of the whole system — the \"Claim\" from Section 3. It is a Java `record` (an "
          "immutable, read-only data holder) with fields like `fieldPath` (which field, e.g. `full_name`), `raw` "
          "(the value exactly as the source had it), `value` (the cleaned value, or `null` if cleaning failed), "
          "`method` (how we got it — used for provenance), `sourceId` (which source), `trust` (how much we "
          "believe it), `observedAt` (a date, used to break ties), and `order` (its position in the manifest, the "
          "final tie-breaker). Because it never changes after creation, the same Claims always produce the same "
          "answer.",
          "A signed witness statement in a courtroom: it records who said what, when, and how reliable the "
          "witness is — and once written, it cannot be altered.")

classcard("model/FieldResolution.java",
          "The winning value for one field, plus its confidence and full evidence.",
          "After the merge step decides a winner for a single scalar field (like `full_name`), it stores the "
          "outcome here: the chosen `value`, its `confidence`, the `winningSource`, the `method`, the list of all "
          "`evidence` Claims that competed, and counts of how many sources `agreed` versus `conflicted`. It is the "
          "\"verdict\" for one field.",
          "A judge's written verdict for a single charge: the decision, how confident the court is, and a record "
          "of every piece of evidence considered.")

classcard("model/CanonicalProfile.java",
          "The fully-assembled, standard internal profile for one candidate.",
          "This is the engine's **internal** picture of the candidate, in a fixed shape that never changes. It "
          "holds the scalar `fields` (each a `FieldResolution`), plus list-style fields that combine rather than "
          "pick one winner — `emails`, `phones`, `skills`, `experience`, `education` — the `provenance` trail, the "
          "`overallConfidence`, and the full `ledger` (kept for auditing). Importantly, this is separate from "
          "whatever the caller finally receives; the projection layer turns it into the requested shape.",
          "The complete, official case file assembled before anyone decides how to present it in a report.")

# --- The adapters package ---
h2("Package: adapters/  — one reader per source format")
body("Adapters are the **translators**. Each one knows how to read a single messy input format and turn it into "
     "clean Claims. The rest of the engine never deals with raw CSV or JSON — only Claims.")

classcard("adapters/Adapter.java",
          "The contract every source-reader must follow.",
          "This is a Java `interface` — a promise that says \"any adapter must provide a `typeName()` and an "
          "`extract(raw)` method that returns a list of `ClaimDraft`s.\" A `ClaimDraft` is a Claim *before* the "
          "pipeline stamps on the shared bookkeeping (source id, trust, timestamp, order). Defining one contract "
          "means adding a brand-new source later is easy: just write a new class that fulfils this promise.",
          "A standard job description: anyone hired as an \"adapter\" must be able to do exactly these tasks, so "
          "they are interchangeable.")

classcard("adapters/Sources.java",
          "All six concrete adapters, plus a registry that maps a source type name to its adapter.",
          "This is the busiest adapter file. It contains the six readers — `RecruiterCsv`, `AtsJson`, `GitHubApi`, "
          "`LinkedIn`, `Resume`, and `RecruiterNotes` — and a `REGISTRY` map so the pipeline can look up the right "
          "reader by name (e.g. `\"recruiter_csv\"`). Two adapters read **structured** data (CSV rows, JSON with "
          "foreign field names); four read **unstructured** data (GitHub/LinkedIn JSON, résumé prose, free-text "
          "notes, using regular expressions to fish out emails, phones, and skills). Each adapter cleans its "
          "values via `Normalize` and guards against bad input so a single broken row degrades to \"no claim\" "
          "instead of crashing. Shared helpers at the bottom handle things like splitting a `\"City, Region, "
          "Country\"` string or parsing a CSV line that contains commas inside quotes.",
          "A team of six specialist translators — one for each language — who all hand their work back in the same "
          "common tongue, plus a receptionist (the registry) who routes each document to the right translator.")

classcard("adapters/Trust.java",
          "The lookup table for how much to believe each source — overall and per field.",
          "Conflict resolution needs a defensible answer to \"who do we believe?\" This class stores the default "
          "trust per source (recruiter CSV 0.90, ATS 0.85, LinkedIn 0.75, GitHub 0.60, résumé 0.55, notes 0.40) "
          "**and** per-field overrides, because a source can be excellent for one field and useless for another. "
          "The standout example: GitHub is trusted **0.95** for skills but **0.00** for phone numbers — it must "
          "never win a phone. The `forField(...)` method returns the effective trust for a given source and field.",
          "A reputation chart for informants: this one is reliable about tech gossip but hopeless about addresses, "
          "so you weight their tips accordingly.")

# --- Top-level engine classes ---
h2("Package: veritas/  — the engine and the command-line tool")

classcard("Json.java",
          "A tiny, self-written JSON reader and writer with zero external libraries.",
          "The project deliberately uses **no third-party code**, so it ships its own small JSON parser and "
          "pretty-printer. `parse(...)` reads JSON text into ordinary Java objects (maps, lists, strings, "
          "numbers); `write(...)` turns them back into neatly-indented JSON text. It preserves insertion order, "
          "which is part of what keeps the output identical on every run. You will see it used by the adapters "
          "(to read JSON sources) and by the CLI (to print the result).",
          "A pocket bilingual dictionary the team carries instead of hiring an outside interpreter — small, but "
          "enough to read and write the one language they need.")

classcard("Normalize.java",
          "Deterministic functions that turn messy values into standard formats — or null.",
          "This class is a toolbox of cleaners: `phoneE164` (→ `+14155550142`), `country` (→ ISO codes like "
          "`US`), `skill` (→ canonical names via an alias table, so `golang`→`Go`, `k8s`→`Kubernetes`), `dateYm` "
          "(→ `YYYY-MM`), `email` (lower-cased and checked), plus `name` and `text` (whitespace tidy-up). The "
          "**load-bearing rule**: every cleaner returns `null` when it cannot be sure, so the system never invents "
          "a value. It also exposes a `REGISTRY` so the config can request a normalizer by name at runtime.",
          "A laundry that returns each garment pressed to the same standard fold — and if something is too damaged "
          "to clean safely, it hands it back untouched rather than ruin it.")

classcard("Merge.java",
          "The heart of the engine: resolve conflicts and compute confidence (the \"reduction\").",
          "This class reduces the whole Evidence Ledger into one `CanonicalProfile`. For **scalar** fields it "
          "picks a single winner using a strict, deterministic order: (1) highest trust, (2) then most recent, "
          "(3) then earliest manifest position, (4) then smallest value alphabetically — the last two guarantee "
          "the same result every run. For **list** fields (emails, phones) it de-duplicates case-insensitively "
          "and unions them. For **skills** it folds together different spellings of the same skill and scores "
          "each by how many sources agree. For **experience/education** it merges records that describe the same "
          "role/school, backfilling blanks. Finally it computes a per-field `confidence` (authority + agreement − "
          "conflict − a penalty for un-cleanable values) and a weighted `overall_confidence` that prioritizes "
          "identity (name, email) over enrichment (skills).",
          "The editor-in-chief: takes every reporter's competing draft, applies clear house rules to choose what "
          "gets printed, and notes in the margin how solid each fact is.")

classcard("Canonical.java",
          "Renders the internal profile into a plain map in the project's default output shape.",
          "`Merge` produces a rich internal object full of engine details. `Canonical.toMap(...)` flattens that "
          "into a clean, ordinary map matching the documented default schema (`candidate_id`, `full_name`, "
          "`emails[]`, `location{...}`, `links{...}`, `skills[...]`, `provenance[...]`, `overall_confidence`, "
          "etc.). It can optionally include or hide confidence and provenance. This is the **clean boundary**: "
          "everything after this point (projection, validation, the CLI) works with this simple map and never "
          "touches the engine's internals.",
          "The press officer who turns the messy internal case file into the standard one-page summary everyone "
          "downstream is allowed to read.")

classcard("Projection.java",
          "The configurable \"twist\": reshape the standard profile using a runtime config — no code changes.",
          "This is the feature that lets one caller ask for the full profile and another ask for a flat contact "
          "card, **without changing the engine**. A config is just data: a list of fields, each with a `path` "
          "(output name), an optional `from` (where to pull it from in the canonical record), a `type` to "
          "validate against, an optional `normalize` step, and whether it is `required`. Projection is a small "
          "interpreter for that config. It understands a mini path language — `emails[0]` (first item), "
          "`location.country` (nested), `skills[].name` (pluck a field from every list element) — selects and "
          "renames fields, applies the requested normalization, type-checks each value, and obeys the "
          "`on_missing` policy (`null`, `omit`, or `error`).",
          "A photographer with one subject and many clients: same person in front of the lens, but each client's "
          "order form (the config) decides the crop, the framing, and what to leave out.")

classcard("Validate.java",
          "A lightweight checker that confirms the default output has the right shape and types.",
          "Before the profile is returned, this class inspects it: is `candidate_id` a string? Are `emails` and "
          "`phones` lists of strings? Is `location` an object with string-or-null parts? Is `overall_confidence` a "
          "number between 0 and 1? It returns a list of human-readable problems — an empty list means \"valid.\" "
          "It checks *structure*, not business rules, because the engine already guarantees the \"unknown → "
          "null\" rule.",
          "The quality-control inspector at the end of the line: doesn't redo the work, just confirms the product "
          "is the right shape before it ships.")

classcard("Pipeline.java",
          "The conductor that runs all seven stages in order.",
          "This class ties everything together: `detect` reads the manifest, `extract` runs each adapter **inside "
          "a sandbox** (a try/catch so a broken source is logged and skipped, never fatal), then it calls "
          "`Merge`, `Canonical`, `Validate`, and optionally `Projection`. It returns a `RunReport` bundling the "
          "candidate id, the canonical record, the final output, any validation errors, a per-source log (what "
          "succeeded, what was skipped and why), and the ledger size. The sandbox in `extract` is what makes the "
          "whole run robust to that corrupt file.",
          "An orchestra conductor: each musician (stage) knows their part, but the conductor decides the order, "
          "keeps time, and makes sure one wrong note doesn't stop the concert.")

classcard("Cli.java",
          "The thin command-line front door — reads arguments, prints or saves the result.",
          "This is what you actually run. It parses command-line flags (`--manifest`, `--config`, `--out`, "
          "`--explain`), loads the optional config, calls `Pipeline.run(...)`, and then either prints the JSON "
          "result to the screen or writes it to a file. With `--explain` it also prints the per-source log to "
          "standard error (which sources gave how many claims, which were skipped, and whether the schema "
          "validated). It deliberately contains almost no logic — all the real work is in the engine.",
          "The reception desk of a building: it greets you, takes your request, sends it to the right department, "
          "and hands back the result — but does none of the actual work itself.")

# --- Tests ---
h2("Package: test/  — proving it works")
classcard("test/Tests.java",
          "A 15-check test suite that runs on a bare JDK (no JUnit needed).",
          "Rather than depend on a testing library, the project ships a tiny home-grown harness. Its 15 checks "
          "cover the normalizers (phone/country/skill/date), end-to-end runs, robustness to the corrupt source, "
          "conflict resolution (the recruiter CSV must win the name), email de-duplication, skill folding, "
          "provenance, **determinism** (two runs must be byte-identical), the projection path language, and all "
          "three `on_missing` policies. The `pipeline_runs_and_validates` check doubles as a \"gold profile\" "
          "test — a fixed input must always produce the known-good output.",
          "A pre-flight checklist a pilot runs every time: a fixed list of yes/no checks that must all pass before "
          "the plane is cleared to fly.")

# ===========================================================================
h1("7 · A worked example: meet Ada Lovelace")
body("The sample data describes one candidate, Ada, across seven sources that disagree on purpose. Here is how "
     "the engine resolves three typical situations.")

h2("7.1 · Resolving the name conflict")
body("Five sources claim a name, with three different spellings. The merge step sorts them by trust (then "
     "recency, then order) and takes the top:")
table(["Value", "Source", "Trust", "Outcome"],
      [["Ada B. Lovelace", "recruiter_csv", "0.90", "**WINNER**"],
       ["Ada Lovelace", "ats_json", "0.85", "conflicts"],
       ["Ada Byron Lovelace", "linkedin", "0.75", "conflicts"],
       ["Ada Lovelace", "github_api", "0.60", "conflicts"],
       ["Ada Lovelace", "resume", "0.55", "conflicts"]])
body("The recruiter CSV wins because it is the most trusted for identity. Crucially, the confidence comes out to "
     "**0.58**, not 1.0 — that is authority `0.90` minus a `0.08` penalty for each of the four disagreeing "
     "sources (`0.90 − 0.08×4 = 0.58`). Veritas flags the uncertainty instead of hiding it. Provenance records "
     "that the value came from `recruiter_csv` via `csv_column:name`. (You can see this exact number in the "
     "output's `field_confidence.full_name`, and the profile's `overall_confidence` is `0.779`.)")

h2("7.2 · Cleaning and de-duplicating phones")
body("The same phone shows up four ways: `(415) 555-0142`, `+1 415 555 0142`, `+1 (415) 555-0142`, and a "
     "different backup `415-555-0199`. Each is normalized to E.164 and duplicates collapse, giving "
     "`[\"+14155550142\", \"+14155550199\"]` — one primary, one backup.")

h2("7.3 · Surviving the corrupt file")
body("`corrupt_ats.json` is intentionally broken JSON. Its adapter throws while parsing, the pipeline's sandbox "
     "catches the error, logs `SKIP (extract failed: ...)`, and the run **continues** on the other six sources. "
     "No value is ever invented from the broken input.")

# ===========================================================================
h1("8 · How to build and run it yourself")
body("The project needs nothing but a JDK 17+. From the project root:")
code([
    "./build.sh                 # compile the engine + tests into ./bin",
    "./test.sh                  # run the 15-check test suite",
    "",
    "# Run with the default profile output:",
    "./run.sh --manifest samples/inputs/manifest.json",
    "",
    "# Run with a custom output config, and show the per-source log:",
    "./run.sh --manifest samples/inputs/manifest.json \\",
    "         --config  samples/configs/compact.json --explain",
    "",
    "# Write the result to a file instead of the screen:",
    "./run.sh --manifest samples/inputs/manifest.json --out out.json",
])
body("The `--explain` flag prints a friendly log to the screen showing how many Claims each source produced, "
     "which were skipped, and whether the output passed validation:")
code([
    "  recruiter_csv    OK   8 claims",
    "  ats_json         OK  16 claims",
    "  ats_json         SKIP (extract failed: ...)   <- corrupt source, run survives",
    "  github_api       OK  10 claims",
    "  ...",
    "  canonical schema: VALID",
])

# ===========================================================================
h1("9 · Where to look next")
body("Now that you know the map, here is a good reading order to explore the actual code:")
numbered("**`model/Claim.java`** — understand the atom first; everything is built on it.")
numbered("**`adapters/Sources.java`** — see how one real source (start with `RecruiterCsv`) becomes Claims.")
numbered("**`Normalize.java`** — see the cleaners that make values comparable.")
numbered("**`Merge.java`** — the most important class: how conflicts are resolved and confidence is scored.")
numbered("**`Projection.java`** — how the same profile becomes different output shapes.")
numbered("**`Pipeline.java`** — finally, see how the conductor calls everything in order.")

closing = doc.add_paragraph()
closing.paragraph_format.space_before = Pt(10)
r = closing.add_run("The one sentence to remember: every source becomes small, traceable Claims; the final "
                    "profile is calculated from those Claims with clear, repeatable rules — so the output is "
                    "trustworthy, explainable, and the same every single time.")
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY

doc.save(OUT)
print("wrote", OUT)
