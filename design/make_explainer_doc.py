#!/usr/bin/env python3
"""Generate the detailed approach explainer (DOCX, ~3 pages).

Covers: the approach, how it solves the problem (with worked examples drawn from
the actual sample run), pros & cons, and further improvements.

Run with the sibling venv that has python-docx:
    ../../eightfold-candidate-transformer/.venv/bin/python make_explainer_doc.py
"""
import os
import re
import subprocess
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "Veritas_Approach_Explained.docx"
ACCENT = RGBColor(0x1F, 0x3B, 0x73)
ACCENT2 = RGBColor(0x2E, 0x5A, 0x88)
GREY = RGBColor(0x55, 0x55, 0x55)
CODEBG = "F2F2F2"

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.55)
    s.left_margin = s.right_margin = Inches(0.65)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(9.5)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.0


def _runs(p, text, size=None):
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not tok:
            continue
        r = p.add_run()
        if size:
            r.font.size = Pt(size)
        if tok.startswith("**") and tok.endswith("**"):
            r.text = tok[2:-2]; r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r.text = tok[1:-1]; r.font.name = "Consolas"; r.font.size = Pt(size or 9)
        else:
            r.text = tok


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ACCENT


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = ACCENT2


def body(text, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    _runs(p, text)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    _runs(p, text)


def code(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.15)
    shade(p, CODEBG)
    for i, line in enumerate(lines):
        r = p.add_run(("" if i == 0 else "\n") + line)
        r.font.name = "Consolas"; r.font.size = Pt(8.5)


def shade(paragraph, fill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def figure(path, caption, width_in=7.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(5)
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        r = c.paragraphs[0].add_run(htext); r.bold = True; r.font.size = Pt(8.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            r = cells[i].paragraphs[0].add_run(str(val)); r.font.size = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ===========================================================================
# Title
# ===========================================================================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(0)
r = t.add_run("Veritas — Approach, Worked Examples, Trade-offs & Roadmap")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = ACCENT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
r = sub.add_run("Multi-Source Candidate Data Transformer · Java (JDK 17, zero dependencies)")
r.font.size = Pt(9); r.font.color.rgb = GREY

# ===========================================================================
h1("1 · The problem, in one breath")
body("Candidate data arrives from many places at once — a recruiter CSV, an ATS JSON blob with its own field "
     "names, a GitHub API response, a LinkedIn profile, résumé prose, free-text recruiter notes. The same person "
     "shows up in several of them with **conflicting** values; any source may be **missing, empty, or malformed**. "
     "Downstream hiring products need exactly **one** clean profile per candidate: fixed fields, normalized "
     "formats, deduplicated, plus a record of **where each value came from and how confident we are**. The hard "
     "rule: *wrong-but-confident is worse than honestly-empty*, because a bad value silently pollutes hiring "
     "decisions.")

# ===========================================================================
h1("2 · The approach — an Evidence Ledger reduced deterministically")
body("**The one idea everything follows from: never merge raw values directly.** Each source is converted into a "
     "list of small, fully-traced assertions called **Claims**. All Claims for a candidate form an **Evidence "
     "Ledger**, and the final profile is a **pure, deterministic reduction** over that ledger. Merging becomes a "
     "function of immutable data — which is what simultaneously makes the system reproducible, explainable, "
     "robust, and easy to scale. The whole flow — heterogeneous **inputs** on the left, the engine in the middle, "
     "the two **output** shapes on the right, and how it meets the three constraints — is shown in Figure 1.")
figure("flow_diagram.png",
       "Figure 1 — High-level design flow: many messy sources → traced Claims in one Evidence Ledger → "
       "deterministic reduction into a canonical record → a config-driven projection into the requested output. "
       "The bottom band shows how each stage satisfies the deterministic, robust, and scalable constraints.")

h2("2.1 · The atom: a Claim")
body("A Claim records one normalized assertion about one canonical field, with its full lineage:")
code([
    'Claim {',
    '  fieldPath  = "full_name"            // canonical target',
    '  raw        = "Ada B. Lovelace"      // exactly as the source had it',
    '  value      = "Ada B. Lovelace"      // normalized (null if it could not be)',
    '  method     = "csv_column:name"      // how we derived it  -> provenance',
    '  sourceId   = "recruiter_csv"',
    '  trust      = 0.90                    // effective trust for THIS field',
    '  observedAt = "2026-06-20"            // recency tie-breaker',
    '  order      = 0                       // manifest position -> final tie-break',
    '}',
])
body("Because Claims keep both `raw` and `value`, the audit trail survives all the way to output, and "
     "normalization failures are visible (a `null` value), never hidden.")

h2("2.2 · The seven-stage pipeline")
body("`detect → extract → normalize → merge → confidence → project → validate`")
bullet("**detect** — read a manifest naming each source (type, path, optional trust + timestamp).")
bullet("**extract** — run the matching adapter inside a **sandbox**; output is a list of claim drafts.")
bullet("**normalize** — happens inside each adapter, so the merge stage only ever sees clean, comparable values.")
bullet("**merge** — reduce the ledger into one canonical record (winner per scalar field; union/dedup per list).")
bullet("**confidence** — score every field from agreement, authority, and conflict.")
bullet("**project** — apply the runtime config to reshape the output (the required twist).")
bullet("**validate** — check the result against the requested schema before returning it.")

h2("2.3 · The trust model — per source AND per field")
body("Conflict resolution needs a defensible notion of *who to believe*. Trust is a per-source base, **overridden "
     "per field** because a source can be authoritative for some fields and useless for others:")
table(["Source", "Base", "Per-field overrides"],
      [["recruiter_csv", "0.90", "—"],
       ["ats_json", "0.85", "—"],
       ["linkedin", "0.75", "headline 0.90, links.linkedin 0.99"],
       ["github_api", "0.60", "skills 0.95, links.github 0.99, phones 0.00"],
       ["resume", "0.55", "experience 0.80, education 0.85, skills 0.70"],
       ["recruiter_notes", "0.40", "phones 0.55, skills 0.45"]])
body("So GitHub is gold for skills but is **forbidden** from ever winning a phone number; the recruiter CSV "
     "outranks LinkedIn on identity. A manifest entry can also override a source's base trust per run.")

h2("2.4 · The confidence model")
body("Scalar-field confidence combines four independent signals, clamped to [0, 1]:")
code([
    'confidence = authority            (winner trust)',
    '           + agreement_bonus      (decaying bonus per corroborating source)',
    '           - conflict_penalty     (0.08 per disagreeing source)',
    '           - normalize_penalty    (0.15 if the winning value never normalized)',
])
body("`overall_confidence` is a weighted mean that deliberately prioritizes **identity and reachability** "
     "(full_name, emails, phones) over enrichment (skills, education).")

h2("2.5 · Canonical record vs. projection layer — two separate types")
body("The engine always builds the **same** internal canonical record. A second, independent layer — the "
     "**projection** — interprets a runtime config (pure data) to reshape that record. The engine never changes "
     "to satisfy a new output shape, which is what makes *“same engine, no code changes”* literally true.")

# ===========================================================================
h1("3 · How it solves the problem — a worked example")
body("Throughout, we use the sample candidate **Ada Lovelace**, whose seven sources disagree on purpose. The "
     "engine turns 63 Claims into one profile with `overall_confidence = 0.779`.")

h2("3.1 · The messy inputs (excerpt)")
table(["Source", "Says the name is…", "Phone", "A few skills"],
      [["recruiter_csv", "Ada B. Lovelace", "(415) 555-0142", "—"],
       ["ats_json", "Ada Lovelace", "+1 415 555 0142", "Python, golang, k8s"],
       ["github_api", "Ada Lovelace", "(none)", "Go, Python, Rust, JavaScript"],
       ["linkedin", "Ada Byron Lovelace", "(none)", "Distributed Systems, Kafka"],
       ["resume", "Ada Lovelace", "+1 (415) 555-0142", "Go, Rust, Spark, Kafka"],
       ["recruiter_notes", "(none)", "415-555-0199", "distributed systems, Go, Kafka"],
       ["corrupt_ats.json", "GARBAGE (broken JSON)", "—", "—"]])

h2("3.2 · Example A — resolving the name conflict")
body("Five sources make a `full_name` Claim with three different values. The reduction sorts them by "
     "(trust ↓, recency ↓, order ↑, value ↑) and takes the head:")
table(["value", "source", "trust", "outcome"],
      [["Ada B. Lovelace", "recruiter_csv", "0.90", "WINNER"],
       ["Ada Lovelace", "ats_json", "0.85", "conflicts"],
       ["Ada Byron Lovelace", "linkedin", "0.75", "conflicts"],
       ["Ada Lovelace", "github_api", "0.60", "conflicts"],
       ["Ada Lovelace", "resume", "0.55", "conflicts"]])
body("Winner = **“Ada B. Lovelace”**. Confidence = authority 0.90 − conflict 0.08×4 = **0.58** — correctly "
     "*lower* than a clean field, signalling the disagreement to downstream consumers. Provenance records "
     "`{field: full_name, source: recruiter_csv, method: csv_column:name}`. The choice is fully reproducible: "
     "even if every trust were equal, the manifest `order` breaks the tie identically every run.")

h2("3.3 · Example B — folding and scoring skills")
body("Skills arrive in many spellings across sources. Normalization canonicalizes aliases "
     "(`golang→Go`, `k8s→Kubernetes`), then a **case-insensitive** fold merges variants — so the free-text "
     "“distributed systems” from notes merges with the ATS “Distributed Systems”. Confidence rises with the "
     "number and trust of corroborating sources:")
table(["skill", "sources", "confidence"],
      [["Go", "ats, github, linkedin, notes, resume (5)", "1.00"],
       ["Distributed Systems", "ats, linkedin, notes (3)", "1.00"],
       ["Apache Kafka", "linkedin, notes, resume (3)", "0.91"],
       ["JavaScript", "github only (1)", "0.95"],
       ["Apache Spark", "resume only (1)", "0.70"]])
body("Note the judgment encoded here: **Go** (5 independent sources) is a near-certainty; **Apache Spark** "
     "(résumé only, trust 0.70) is kept but flagged as weaker — we never *drop* a niche skill, we just score it "
     "honestly.")

h2("3.4 · Example C — normalization & dedup")
body("Phones appear as `(415) 555-0142`, `+1 415 555 0142`, `+1 (415) 555-0142`, `415-555-0199`. Each is "
     "normalized to **E.164** and deduped, yielding `[\"+14155550142\", \"+14155550199\"]` — the four spellings of "
     "the primary collapse to one, and the backup mobile from the notes is preserved. Emails are lower-cased and "
     "deduped: `Ada.Lovelace@example.com` and `ada.lovelace@example.com` become one entry, while "
     "`ALOVELACE@oldmail.com` is kept as a second address. Country `United States`/`USA` → **US** "
     "(ISO-3166), confidence 1.0 (every source agrees).")

h2("3.5 · Example D — robustness to a garbage source")
body("`corrupt_ats.json` is intentionally invalid JSON. Its adapter throws while parsing; the pipeline's "
     "per-source sandbox catches it, records `SKIP (extract failed: …)`, and **the run continues** on the other "
     "six sources. No field is invented from the broken input — exactly the *“honestly-empty beats "
     "wrong-but-confident”* rule in action.")

h2("3.6 · Example E — the configurable twist")
body("A caller wants a flat contact card, not the full profile. They send a config — **no engine change**:")
code([
    '{ "fields": [',
    '   { "path": "full_name",     "type": "string", "required": true },',
    '   { "path": "primary_email", "from": "emails[0]", "type": "string", "required": true },',
    '   { "path": "phone",         "from": "phones[0]", "type": "string", "normalize": "E164" },',
    '   { "path": "skills",        "from": "skills[].name", "type": "string[]" } ],',
    '  "on_missing": "omit" }',
])
body("The projection layer remaps `emails[0] → primary_email`, plucks `skills[].name` into a flat list, "
     "type-validates each value, and omits anything missing. Result:")
code([
    '{ "full_name": "Ada B. Lovelace",',
    '  "primary_email": "ada.lovelace@example.com",',
    '  "phone": "+14155550142",',
    '  "skills": ["Apache Kafka","Apache Spark","Distributed Systems","Go", ...],',
    '  "overall_confidence": 0.779 }',
])
body("Switch `on_missing` to `error` and a missing required field raises instead; switch a type and an invalid "
     "value is rejected at the boundary. Same canonical record, three different contracts.")

h2("3.7 · …which maps cleanly to the three constraints")
bullet("**Deterministic & explainable** — pure reduction over immutable Claims; total ordering of tie-breaks; "
       "every output field carries provenance + confidence.")
bullet("**Robust** — per-source sandboxes; normalizers return `null` rather than guess; validation gate before return.")
bullet("**Scale** — each candidate is independent and O(claims); adapters are stateless and parallelizable; no "
       "cross-candidate state.")

# ===========================================================================
h1("4 · Pros of this approach")
bullet("**Explainability is free, not bolted on.** Provenance and confidence fall directly out of the ledger; "
       "you can always answer “why does the profile say X?”.")
bullet("**Determinism by construction.** Same inputs → byte-identical output, because the result is a pure "
       "function of immutable Claims with a total tie-break order. Trivial to gold-test.")
bullet("**Policy lives in one place.** Trust and conflict rules sit in the merge stage, not smeared across "
       "adapters — easy to read, defend, and change.")
bullet("**Open to extension, closed to modification.** A new source = a new adapter implementing one interface; "
       "a new output shape = a new config, no code. Adapters never know about each other.")
bullet("**Fails safe.** The default outcome of any uncertainty is `null` + low confidence, never a fabricated value.")
bullet("**Zero-dependency & portable.** Bare JDK, no build tool or network required — easy to run and audit.")

h1("5 · Cons & honest limitations")
bullet("**Trust weights are hand-tuned heuristics.** They are sensible defaults, not learned from outcomes; a "
       "miscalibrated weight can pick the wrong winner on an unusual candidate.")
bullet("**Confidence is a designed score, not a probability.** The numbers are useful for ranking and "
       "thresholds, but they are not calibrated likelihoods.")
bullet("**Single-record identity assumption.** One candidate per manifest; the approach does not yet do "
       "cross-candidate entity resolution (fuzzy matching the same person across files).")
bullet("**Deterministic parsing under-extracts.** The résumé reader is regex + section headers (for "
       "reproducibility), so it deliberately misses prose it cannot parse confidently.")
bullet("**Ledger memory cost.** Keeping every Claim is great for auditing but is more memory than a "
       "merge-in-place approach; for very wide fan-in this would need streaming/spilling.")
bullet("**Normalization tables are finite.** Country and skill-alias maps are curated; an unknown alias passes "
       "through verbatim (safe, but not canonical).")

h1("6 · How it can be improved")
bullet("**Learn trust from feedback.** Replace static weights with per-(source, field) reliability learned from "
       "recruiter corrections or downstream outcomes — turning the heuristic into a calibrated prior.")
bullet("**Calibrated, probabilistic confidence.** Fit the score to observed correctness (e.g. isotonic/Platt "
       "calibration) so `0.8` genuinely means 80% right.")
bullet("**Cross-candidate identity resolution.** Add a blocking + fuzzy-match stage (name/email/phone "
       "similarity) so the same person across sources merges into one record.")
bullet("**Stronger extractors behind the same contract.** Swap regex résumé parsing for an LLM/NER extractor "
       "that still emits Claims — the ledger and merge stay unchanged, and we keep determinism via cached, "
       "versioned extractions.")
bullet("**Real connectors.** Replace the captured GitHub JSON with a live `HttpClient` fetch (with caching for "
       "reproducibility); add LinkedIn/ATS API adapters — again, only new adapters, no engine change.")
bullet("**Schema-driven validation & config.** Express the default schema and configs as JSON-Schema and "
       "validate against it; add a tiny UI over the existing config for non-engineers.")
bullet("**Observability.** Emit per-run metrics (claims per source, conflict rate, fields left null) to catch a "
       "drifting source before it pollutes profiles.")

closing = doc.add_paragraph()
closing.paragraph_format.space_before = Pt(8)
r = closing.add_run("In short: model every assertion as traceable evidence, reduce it with an explainable, "
                    "deterministic policy, and keep the output shape a matter of configuration — so the system is "
                    "trustworthy first, and flexible without becoming fragile.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY

doc.save(OUT)
print("wrote", OUT)
